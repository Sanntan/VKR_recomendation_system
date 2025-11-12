from telegram import Update, InlineKeyboardMarkup, InlineKeyboardButton
from telegram.ext import ContextTypes
from datetime import datetime, date
from uuid import UUID
from typing import Any, Mapping, Dict
import io

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.bot.services.api_client import api_client, APIClientError
from src.bot.middlewares.auth_middleware import auth_required

def _parse_date(value: Any) -> date | None:
    if not value:
        return None
    if isinstance(value, date):
        return value
    if isinstance(value, str):
        try:
            return datetime.fromisoformat(value).date()
        except ValueError:
            try:
                return datetime.strptime(value, "%Y-%m-%d").date()
            except ValueError:
                return None
    return None


def _get_value(event: Mapping[str, Any] | Any, attr: str) -> Any:
    if isinstance(event, Mapping):
        return event.get(attr)
    return getattr(event, attr, None)


def format_event_card(event: Mapping[str, Any] | Any) -> str:
    """Форматирует карточку мероприятия, поддерживая словари и ORM объекты."""
    start_raw = _get_value(event, "start_date")
    end_raw = _get_value(event, "end_date")
    start_date = _parse_date(start_raw)
    end_date = _parse_date(end_raw)

    start_str = start_date.strftime('%d.%m.%Y') if start_date else 'Не указана'
    end_str = end_date.strftime('%d.%m.%Y') if end_date else ''

    date_str = start_str
    if end_str and start_str != end_str:
        date_str = f"{start_str} - {end_str}"

    title = _get_value(event, "title") or "Без названия"
    short_description = _get_value(event, "short_description")
    format_value = _get_value(event, "format")
    link = _get_value(event, "link")
    likes_count = _get_value(event, "likes_count") or 0
    dislikes_count = _get_value(event, "dislikes_count") or 0

    text = f"🎯 *{title}*\n\n"

    if short_description:
        text += f"{short_description}\n\n"

    text += f"📅 Дата: {date_str}\n"

    if format_value:
        text += f"🎯 Формат: {format_value}\n"

    if link:
        text += f"🔗 [Регистрация]({link})\n"

    text += f"👍 {likes_count} 👎 {dislikes_count}"

    return text

def get_recommendation_buttons(event_id: str) -> InlineKeyboardMarkup:
    """Создает кнопки для взаимодействия с рекомендацией."""
    keyboard = [
        [
            InlineKeyboardButton("👍 Интересно", callback_data=f"like_{event_id}"),
            InlineKeyboardButton("👎 Не интересно", callback_data=f"dislike_{event_id}")
        ],
        [
            InlineKeyboardButton("🔄 Показать другие", callback_data="show_other_events"),
            InlineKeyboardButton("🔙 Назад", callback_data="back_to_menu")
        ]
    ]
    return InlineKeyboardMarkup(keyboard)

@auth_required
async def show_recommendations(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Показывает рекомендации пользователю."""
    student = context.user_data.get('student')
    if not student:
        await update.callback_query.edit_message_text(
            "❌ Ошибка: пользователь не найден. Попробуйте авторизоваться заново."
        )
        return

    student_id = student.get("id")
    if not student_id:
        await update.callback_query.edit_message_text(
            "❌ Ошибка: не найден идентификатор студента."
        )
        return

    try:
        student_uuid = UUID(student_id)
    except (ValueError, TypeError):
        await update.callback_query.edit_message_text(
            "Ошибка идентификации студента. Попробуйте авторизоваться заново."
        )
        return

    try:
        recommendations = await api_client.get_recommendations(student_uuid, limit=10)
    except APIClientError:
        await update.callback_query.edit_message_text(
            "Не удалось загрузить рекомендации. Попробуйте позже.",
            reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("🔙 Назад", callback_data="back_to_menu")]])
        )
        return

    if not recommendations:
        await update.callback_query.edit_message_text(
            "Пока нет подходящих мероприятий для рекомендаций.\n"
            "Попробуйте воспользоваться поиском мероприятий!",
            reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("🔍 Поиск мероприятий", callback_data="event_search")]])
        )
        return

    context.user_data['current_recommendations'] = recommendations
    context.user_data['current_recommendation_index'] = 0

    event_ids = [rec.get("event_id") for rec in recommendations if rec.get("event_id")]
    events_cache: Dict[str, Any] = {}
    if event_ids:
        try:
            bulk_response = await api_client.get_events_bulk(event_ids)
            for event in bulk_response.get("events", []):
                events_cache[str(event["id"])] = event
        except APIClientError:
            events_cache = {}

    context.user_data['recommendations_events'] = events_cache

    first_rec = recommendations[0]
    event_id = first_rec.get("event_id")
    event = events_cache.get(str(event_id))

    if not event:
        if event_id:
            try:
                event = await api_client.get_event(UUID(event_id))
            except APIClientError:
                event = None
            if event:
                events_cache[str(event["id"])] = event

    if not event:
        await update.callback_query.edit_message_text(
            "Ошибка загрузки мероприятия. Попробуйте позже.",
            reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("🔙 Назад", callback_data="back_to_menu")]])
        )
        return

    await update.callback_query.edit_message_text(
        format_event_card(event),
        reply_markup=get_recommendation_buttons(str(event["id"])),
        parse_mode='Markdown',
        disable_web_page_preview=True
    )

@auth_required
async def handle_recommendation_feedback(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Обрабатывает feedback по рекомендациям."""
    query = update.callback_query
    await query.answer()

    action, event_id_str = query.data.split('_')
    event_uuid = UUID(event_id_str)

    try:
        if action == 'like':
            updated_event = await api_client.like_event(event_uuid)
            event_cache = context.user_data.get('recommendations_events', {})
            if isinstance(event_cache, dict) and updated_event:
                event_cache[str(updated_event["id"])] = updated_event
                context.user_data['recommendations_events'] = event_cache
            await query.answer("Спасибо! Учтем ваши предпочтения 👍")
        elif action == 'dislike':
            await api_client.dislike_event(event_uuid)
            await show_next_recommendation(update, context)
            return
    except APIClientError:
        await query.answer("Ошибка при обработке запроса. Попробуйте позже.", show_alert=True)
        return

    event_cache = context.user_data.get('recommendations_events', {})
    event = {}
    if isinstance(event_cache, dict):
        event = event_cache.get(str(event_uuid), {})

    if not event:
        try:
            event = await api_client.get_event(event_uuid)
        except APIClientError:
            event = None
        if event and isinstance(event_cache, dict):
            event_cache[str(event["id"])] = event
            context.user_data['recommendations_events'] = event_cache

    if event:
        await query.edit_message_text(
            format_event_card(event),
            reply_markup=get_recommendation_buttons(str(event["id"])),
            parse_mode='Markdown',
            disable_web_page_preview=True
        )

@auth_required
async def show_next_recommendation(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Показывает следующую рекомендацию."""
    query = update.callback_query
    await query.answer()

    recommendations = context.user_data.get('current_recommendations', [])
    current_index = context.user_data.get('current_recommendation_index', 0)

    if not recommendations:
        await show_recommendations(update, context)
        return

    # Переходим к следующему мероприятию
    current_index = (current_index + 1) % len(recommendations)
    context.user_data['current_recommendation_index'] = current_index

    event_id = recommendations[current_index].get("event_id")
    if not event_id:
        await show_recommendations(update, context)
        return

    events_cache = context.user_data.get('recommendations_events', {})
    event = {}
    if isinstance(events_cache, dict):
        event = events_cache.get(str(event_id), {})

    if not event:
        try:
            event = await api_client.get_event(UUID(event_id))
        except APIClientError:
            event = None
        if not event:
            await show_recommendations(update, context)
            return
        if isinstance(events_cache, dict):
            events_cache[str(event["id"])] = event
            context.user_data['recommendations_events'] = events_cache

    await query.edit_message_text(
        format_event_card(event),
        reply_markup=get_recommendation_buttons(str(event["id"])),
        parse_mode='Markdown',
        disable_web_page_preview=True
    )


def create_recommendations_docx(recommendations: list[Dict[str, Any]], events: Dict[str, Dict[str, Any]]) -> io.BytesIO:
    """Создает DOCX файл с рекомендациями, отсортированными по score (от самых близких)."""
    doc = Document()
    
    # Заголовок документа
    title = doc.add_heading('Мои рекомендации', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Добавляем дату создания
    date_para = doc.add_paragraph(f'Дата создания: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para_format = date_para.runs[0].font
    date_para_format.size = Pt(10)
    date_para_format.italic = True
    
    doc.add_paragraph()  # Пустая строка
    
    # Сортируем рекомендации по score (от большего к меньшему)
    sorted_recommendations = sorted(
        recommendations,
        key=lambda x: x.get("score", 0.0),
        reverse=True
    )
    
    if not sorted_recommendations:
        doc.add_paragraph("Пока нет рекомендаций.")
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        return doc_buffer
    
    # Добавляем каждую рекомендацию
    for idx, rec in enumerate(sorted_recommendations, 1):
        event_id = str(rec.get("event_id", ""))
        event = events.get(event_id)
        
        if not event:
            continue
        
        # Заголовок мероприятия
        event_title = doc.add_heading(f'{idx}. {event.get("title", "Без названия")}', level=1)
        event_title_format = event_title.runs[0].font
        event_title_format.size = Pt(14)
        event_title_format.bold = True
        
        # Описание
        if event.get("short_description"):
            desc_para = doc.add_paragraph(event["short_description"])
            desc_format = desc_para.runs[0].font
            desc_format.size = Pt(11)
        
        # Информация о мероприятии
        info_para = doc.add_paragraph()
        
        # Дата
        start_raw = event.get("start_date")
        end_raw = event.get("end_date")
        start_date = _parse_date(start_raw)
        end_date = _parse_date(end_raw)
        
        if start_date:
            start_str = start_date.strftime('%d.%m.%Y')
            end_str = end_date.strftime('%d.%m.%Y') if end_date else ''
            date_str = start_str
            if end_str and start_str != end_str:
                date_str = f"{start_str} - {end_str}"
            info_para.add_run("📅 Дата: ").bold = True
            info_para.add_run(date_str)
            info_para.add_run("\n")
        
        # Формат
        if event.get("format"):
            info_para.add_run("🎯 Формат: ").bold = True
            info_para.add_run(event["format"])
            info_para.add_run("\n")
        
        # Ссылка
        if event.get("link"):
            info_para.add_run("🔗 Ссылка: ").bold = True
            info_para.add_run(event["link"])
            info_para.add_run("\n")
        
        # Оценка релевантности
        score = rec.get("score")
        if score is not None:
            info_para.add_run("⭐ Оценка релевантности: ").bold = True
            info_para.add_run(f"{score:.2f}")
            info_para.add_run("\n")
        
        # Лайки/дизлайки
        likes = event.get("likes_count", 0)
        dislikes = event.get("dislikes_count", 0)
        info_para.add_run(f"👍 {likes}  👎 {dislikes}")
        
        # Разделитель между мероприятиями
        if idx < len(sorted_recommendations):
            doc.add_paragraph("─" * 50)
            doc.add_paragraph()
    
    # Сохраняем документ в BytesIO
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer


@auth_required
async def export_recommendations(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Выгружает все рекомендации пользователя в DOCX файл."""
    query = update.callback_query
    if query:
        await query.answer()
    
    student = context.user_data.get('student')
    if not student:
        error_text = "❌ Ошибка: пользователь не найден. Попробуйте авторизоваться заново."
        if query:
            await query.edit_message_text(
                error_text,
                reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("🔙 Назад", callback_data="back_to_menu")]])
            )
        else:
            await update.message.reply_text(error_text)
        return
    
    student_id = student.get("id")
    if not student_id:
        error_text = "❌ Ошибка: не найден идентификатор студента."
        if query:
            await query.edit_message_text(
                error_text,
                reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("🔙 Назад", callback_data="back_to_menu")]])
            )
        else:
            await update.message.reply_text(error_text)
        return
    
    try:
        student_uuid = UUID(student_id)
    except (ValueError, TypeError):
        error_text = "Ошибка идентификации студента. Попробуйте авторизоваться заново."
        if query:
            await query.edit_message_text(
                error_text,
                reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("🔙 Назад", callback_data="back_to_menu")]])
            )
        else:
            await update.message.reply_text(error_text)
        return
    
    # Показываем сообщение о начале генерации
    loading_text = "⏳ Формирую файл с рекомендациями..."
    if query:
        await query.edit_message_text(loading_text)
    else:
        loading_msg = await update.message.reply_text(loading_text)
    
    try:
        # Получаем все рекомендации (большой лимит)
        recommendations = await api_client.get_recommendations(student_uuid, limit=1000)
    except APIClientError:
        error_text = "Не удалось загрузить рекомендации. Попробуйте позже."
        if query:
            await query.edit_message_text(
                error_text,
                reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("🔙 Назад", callback_data="back_to_menu")]])
            )
        else:
            await update.message.reply_text(error_text)
        return
    
    if not recommendations:
        error_text = (
            "Пока нет подходящих мероприятий для рекомендаций.\n"
            "Попробуйте воспользоваться поиском мероприятий!"
        )
        if query:
            await query.edit_message_text(
                error_text,
                reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("🔍 Поиск мероприятий", callback_data="event_search")]])
            )
        else:
            await update.message.reply_text(error_text)
        return
    
    # Получаем информацию о событиях
    event_ids = [rec.get("event_id") for rec in recommendations if rec.get("event_id")]
    events: Dict[str, Dict[str, Any]] = {}
    
    if event_ids:
        try:
            bulk_response = await api_client.get_events_bulk(event_ids)
            for event in bulk_response.get("events", []):
                events[str(event["id"])] = event
        except APIClientError:
            # Если bulk не сработал, получаем по одному
            for event_id in event_ids:
                try:
                    event = await api_client.get_event(UUID(event_id))
                    if event:
                        events[str(event["id"])] = event
                except APIClientError:
                    continue
    
    # Создаем DOCX файл
    try:
        docx_buffer = create_recommendations_docx(recommendations, events)
    except Exception as e:
        error_text = f"Ошибка при создании файла: {str(e)}"
        if query:
            await query.edit_message_text(
                error_text,
                reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("🔙 Назад", callback_data="back_to_menu")]])
            )
        else:
            await update.message.reply_text(error_text)
        return
    
    # Отправляем файл
    filename = f"recommendations_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    try:
        if query:
            # Отправляем файл
            await query.message.reply_document(
                document=docx_buffer,
                filename=filename,
                caption=f"📄 Ваши рекомендации ({len(recommendations)} мероприятий)"
            )
            # Удаляем сообщение с загрузкой
            try:
                await query.message.delete()
            except Exception:
                # Если не удалось удалить, просто редактируем на финальное сообщение
                await query.edit_message_text(
                    f"✅ Файл успешно отправлен!\n📄 Ваши рекомендации ({len(recommendations)} мероприятий)"
                )
        else:
            await update.message.reply_document(
                document=docx_buffer,
                filename=filename,
                caption=f"📄 Ваши рекомендации ({len(recommendations)} мероприятий)"
            )
    except Exception as e:
        error_text = f"Ошибка при отправке файла: {str(e)}"
        if query:
            await query.edit_message_text(
                error_text,
                reply_markup=InlineKeyboardMarkup([[InlineKeyboardButton("🔙 Назад", callback_data="back_to_menu")]])
            )
        else:
            await update.message.reply_text(error_text)
        return